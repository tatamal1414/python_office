from docx import Document

document = Document()

document.add_heading('簡単なWordドキュメントのタイトル', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

document.save('sample.docx')

document.add_picture('character_main.png')

doc = Document("sample.docx")
num = 0
for para in doc.paragraphs:
    num += len(para.text)
    print(num, para.text)
    print(len(para.text))
    print(num)

document.add_paragraph(f'文字数{num}')



document.save('sample_answer.docx')